#!/usr/bin/env python3
"""Build the project manual DOCX from its Markdown source.

This script intentionally depends only on python-docx so it can run with the
bundled Codex document runtime.  The output follows the compact reference guide
style used for a dense, procedural course manual.
"""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
CODE_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, ascii_name="Calibri", east_asia_name="Microsoft YaHei", size=None):
    run.font.name = ascii_name
    if size is not None:
        run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), ascii_name)
    fonts.set(qn("w:hAnsi"), ascii_name)
    fonts.set(qn("w:eastAsia"), east_asia_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn("w:" + side))
        if element is None:
            element = OxmlElement("w:" + side)
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            repeat_header = OxmlElement("w:tblHeader")
            repeat_header.set(qn("w:val"), "true")
            tr_pr.append(repeat_header)


def create_decimal_abstract_numbering(document):
    numbering = document.part.numbering_part.element
    existing = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing) + 1 if existing else 0

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), "71C17E07")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    level.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    ppr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    abstract.append(level)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)
    return abstract_id


def create_numbering_instance(document, abstract_id):
    numbering = document.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing) + 1 if existing else 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_ref)


def choose_widths(rows):
    count = max(len(row) for row in rows)
    weights = []
    for column in range(count):
        lengths = [len(row[column]) if column < len(row) else 0 for row in rows]
        weights.append(max(5, min(max(lengths), 48)))
    minimum = 720
    available = CONTENT_WIDTH_DXA - minimum * count
    total = sum(weights) or count
    widths = [minimum + int(available * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_inline(paragraph, text, bold_default=False):
    token_re = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            run.bold = bold_default
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            rpr = run._element.get_or_add_rPr()
            rpr.rFonts.set(qn("w:ascii"), "Consolas")
            rpr.rFonts.set(qn("w:hAnsi"), "Consolas")
            rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()
            run = paragraph.add_run(label)
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True
            run.bold = bold_default
            set_run_font(run)
            if url.startswith("http"):
                tooltip = paragraph.add_run(" <{}>".format(url))
                tooltip.font.size = Pt(8)
                tooltip.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                set_run_font(tooltip, size=8)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.bold = bold_default
        set_run_font(run)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(16)
    subtitle.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = document.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.extend([begin, instr, separate, value, end])
    paragraph._p.append(field_run)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("CITEL-T-007  |  XML 协议 UDP 发生器")
    set_run_font(run, size=8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_field(footer)


def split_table_row(line):
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    cells = []
    current = []
    inside_code = False
    escaped = False
    for character in stripped:
        if escaped:
            current.append(character)
            escaped = False
        elif character == "\\":
            current.append(character)
            escaped = True
        elif character == "`":
            current.append(character)
            inside_code = not inside_code
        elif character == "|" and not inside_code:
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(character)
    cells.append("".join(current).strip())
    return cells


def is_separator_row(cells):
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def add_table(document, rows):
    if len(rows) >= 2 and is_separator_row(rows[1]):
        rows = [rows[0]] + rows[2:]
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    widths = choose_widths(rows)
    for row_index, source_row in enumerate(rows):
        for column_index in range(column_count):
            text = source_row[column_index] if column_index < len(source_row) else ""
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline(paragraph, text, bold_default=(row_index == 0))
            if row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
    set_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def add_code_block(document, lines):
    paragraph = document.add_paragraph(style="Code Block")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, ascii_name="Consolas", size=8.5)
        if index != len(lines) - 1:
            run.add_break()
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_FILL)
    ppr.append(shd)


def add_callout(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    add_inline(paragraph, text)
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    ppr.append(shd)


def build(markdown_path, output_path):
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_styles(document)
    configure_document(document)
    document.core_properties.title = "CITEL-T-007 XML 协议 UDP 发生器完整项目说明书"
    document.core_properties.subject = "Qt 5.9.7 课程项目操作、验收与复现手册"
    document.core_properties.author = "CITEL-T-007 项目组"
    decimal_abstract_id = create_decimal_abstract_numbering(document)

    index = 0
    first_title = True
    active_numbering_id = None
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered is None:
            active_numbering_id = None
        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(document, code_lines)
            index += 1
            continue

        if stripped.startswith("|"):
            table_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(split_table_row(lines[index]))
                index += 1
            add_table(document, table_rows)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and first_title:
                paragraph = document.add_paragraph(style="Title")
                add_inline(paragraph, text)
                first_title = False
            elif level == 2 and text == "完整项目说明书":
                paragraph = document.add_paragraph(style="Subtitle")
                add_inline(paragraph, text)
            else:
                style = "Heading {}".format(level)
                paragraph = document.add_paragraph(style=style)
                add_inline(paragraph, text)
            index += 1
            continue

        if stripped.startswith(">"):
            add_callout(document, stripped[1:].strip())
            index += 1
            continue

        checklist = re.match(r"^-\s+\[([ xX])\]\s+(.+)$", stripped)
        if checklist:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.187)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_after = Pt(4)
            add_inline(paragraph, ("☒ " if checklist.group(1).lower() == "x" else "☐ ") + checklist.group(2))
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue

        if numbered:
            if active_numbering_id is None:
                active_numbering_id = create_numbering_instance(document, decimal_abstract_id)
            paragraph = document.add_paragraph(style="List Number")
            apply_numbering(paragraph, active_numbering_id)
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline(paragraph, stripped)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.input.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
